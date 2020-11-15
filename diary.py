from tkinter import Tk, Label, LabelFrame, W, messagebox, Text, END, Listbox, ANCHOR, WORD, \
    Scrollbar, Y, Frame, HORIZONTAL, VERTICAL, RIGHT, ttk, StringVar, Toplevel, filedialog
import sqlite3
from datetime import datetime
import webbrowser
from PIL import ImageTk, Image
from docx import Document
import os
import shutil


def main():
    # window diary
    def main_window():
        def _on_key_release(event):
            ctrl = (event.state & 0x4) != 0
            if event.keycode == 88 and ctrl and event.keysym.lower() != "x":
                event.widget.event_generate("<<Cut>>")

            if event.keycode == 86 and ctrl and event.keysym.lower() != "v":
                event.widget.event_generate("<<Paste>>")

            if event.keycode == 67 and ctrl and event.keysym.lower() != "c":
                event.widget.event_generate("<<Copy>>")

            if event.keycode == 65 and ctrl and event.keysym.lower() != "a":
                event.widget.event_generate("<<SelectAll>>")

        def save():
            nonlocal count_record, picture_symbol, picture_n, full_path_picture
            time = datetime.now().strftime("%d.%m.%Y - %H:%M:%S")
            if full_path_picture != '':
                old_path = full_path_picture
                if full_path_picture.rfind('logs_ps\\') != -1:
                    full_path_picture = full_path_picture[
                                        :full_path_picture.rfind('logs_ps\\')] + 'pictures\\' + picture_n
                    shutil.copyfile(old_path, full_path_picture)
                    os.remove(old_path)
                elif full_path_picture.rfind('pictures\\') != -1:
                    picture_n = str(picture_n[0])
                    full_path_picture = full_path_picture[
                                        :full_path_picture.rfind('pictures\\')] + 'pictures\\copy_' + picture_n[
                                        :-4] + datetime.now().strftime("%d-%m-%Y %H.%M.%S") + '.jpg'
                    shutil.copyfile(old_path, full_path_picture)
                    picture_n = os.path.basename(full_path_picture)
            try:
                cursor.execute(
                    f"INSERT INTO {table_name} (login, password, date, name, text, picture_name, "
                    f"picture_path, picture_symbol) VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
                    (user_login, user_password, time, name.get(), text.get('1.0', END), picture_n, full_path_picture,
                     picture_symbol))
            except sqlite3.InterfaceError:
                cursor.execute(
                    f"INSERT INTO {table_name} (login, password, date, name, text, picture_name, "
                    f"picture_path, picture_symbol) VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
                    (user_login, user_password, time, name.get(), text.get('1.0', END), str(picture_n[0]),
                     full_path_picture, picture_symbol))
            db.commit()
            list_records.delete(0, END)
            count_record = 0
            for item1 in cursor.execute(f'SELECT id, date, picture_symbol, name FROM {table_name}'):
                if item1[2] == '':
                    item_list = list(item1)
                    item_list.remove('')
                    item1 = item_list
                    list_records.insert(0, item1)
                    count_record += 1
                else:
                    list_records.insert(0, item1)
                    count_record += 1
            journal_lab = Label(diary, text=f"Журнал записей (всего: {count_record})", bg='white',
                                font=('Comic Sans MS', 11))
            journal_lab.place(x=26, y=125)
            delete_button['state'] = 'normal'
            open_button['state'] = 'normal'
            picture_n = ''
            full_path_picture = ''
            picture_symbol = ''
            attached_picture_but['state'] = 'disabled'
            name_label_picture_main['text'] = picture_n
            name.delete(0, END)
            name.insert(0, 'no_name')
            name['foreground'] = 'gray55'
            text.delete(1.0, END)

        def setting():
            nonlocal last_size, last_font
            settings = Tk()
            settings.title('Settings')
            x_value = (settings.winfo_screenwidth() - 375) / 2
            y_value = (settings.winfo_screenheight() - 160) / 2
            settings.wm_geometry("+%d+%d" % (x_value, y_value))
            settings.geometry('375x160')
            settings.resizable(False, False)
            settings['bg'] = 'gray90'

            keys_log = Label(settings, text='Сочетание клавиш для быстрого входа: ctrl +        (англ.буква)',
                             bg='gray90')
            keys_log.place(relx=0.04, rely=0.15)

            key_two = ttk.Entry(settings, width=2)
            key_two.place(relx=0.725, rely=0.144)
            key_two.insert(1, install_char)

            # Label font
            ttk.Label(settings, text="Шрифт:", background='gray90').place(relx=0.043, rely=0.32)

            n = StringVar()
            shrift_choosen = ttk.Combobox(settings, width=22, textvariable=n, state='readonly')

            # Adding combobox drop down list
            shrift_choosen['values'] = ('Comic Sans MS', 'Segoe Script', 'Cambria')

            shrift_choosen.place(relx=0.17, rely=0.31)

            # Shows comic sans MS as a default value
            for item2, numb in enumerate(shrift_choosen['values']):
                if numb == last_font:
                    shrift_choosen.current(item2)

            def set_font(_):
                nonlocal last_font
                last_font = shrift_choosen.get()

            shrift_choosen.bind("<<ComboboxSelected>>", set_font)
            # Label font size
            ttk.Label(settings, text="Размер шрифта:", background='gray90').place(relx=0.043, rely=0.49)
            n2 = StringVar()
            size_choosen = ttk.Combobox(settings, width=3, textvariable=n2, state='readonly')
            # Adding combobox drop down list
            size_choosen['values'] = ('10', '12', '14')
            size_choosen.place(relx=0.30, rely=0.48)
            # Shows 12 as a default value
            for item3, numb in enumerate(size_choosen['values']):
                if numb == last_size:
                    size_choosen.current(item3)

            def set_port(_):
                nonlocal last_size
                last_size = size_choosen.get()

            size_choosen.bind("<<ComboboxSelected>>", set_port)

            def add_keys():
                nonlocal enter
                enter = key_two.get()
                cursor.execute("INSERT INTO last_user (login, password, char) VALUES (?, ?, ?)",
                               (user_login, user_password, enter))
                db.commit()
                if last_font == 'Comic Sans MS' and last_size == '10':
                    text.configure(font=(last_font, int(last_size)), height=15, width=72)
                elif last_font == 'Comic Sans MS' and last_size == '12':
                    text.configure(font=(last_font, int(last_size)), height=12, width=58)
                elif last_font == 'Comic Sans MS' and last_size == '14':
                    text.configure(font=(last_font, int(last_size)), height=11, width=48)
                elif last_font == 'Segoe Script' and last_size == '10':
                    text.configure(font=(last_font, int(last_size)), height=14, width=58)
                elif last_font == 'Segoe Script' and last_size == '12':
                    text.configure(font=(last_font, int(last_size)), height=11, width=48)
                elif last_font == 'Segoe Script' and last_size == '14':
                    text.configure(font=(last_font, int(last_size)), height=9, width=41)
                elif last_font == 'Cambria' and last_size == '10':
                    text.configure(font=(last_font, int(last_size)), height=19, width=83)
                elif last_font == 'Cambria' and last_size == '12':
                    text.configure(font=(last_font, int(last_size)), height=15, width=65)
                elif last_font == 'Cambria' and last_size == '14':
                    text.configure(font=(last_font, int(last_size)), height=13, width=53)
                settings.destroy()

            style_keys = ttk.Style()
            style_keys.configure('keys.TButton', font=('Comic Sans MS', 8))
            keys_button = ttk.Button(settings, text='Применить', command=add_keys)

            keys_button.place(relx=0.75, rely=0.74)

        def delete_trash():
            nonlocal picture_n, picture_symbol
            try:
                if full_path_picture.rfind('logs_ps') != -1:
                    os.remove(full_path_picture)
                    picture_n = ''
                    name_label_picture_main['text'] = picture_n
                    attached_picture_but['state'] = 'disabled'
                    picture_symbol = ''
                elif full_path_picture.rfind('pictures') != -1:
                    picture_n = ''
                    name_label_picture_main['text'] = picture_n
                    attached_picture_but['state'] = 'disabled'
                    picture_symbol = ''
            except FileNotFoundError:
                pass

        def delete():
            nonlocal count_record, full_path_picture
            try:
                date_rec = list_records.get(ANCHOR)
                result = messagebox.askyesno('My diary', f'Вы действительно хотите удалить запись "{date_rec[0]}"?')
                if result:
                    delete_trash()
                    for item_delete in cursor.execute(
                            f'SELECT picture_path FROM {table_name} WHERE id = {date_rec[0]}'):
                        full_path_picture = str(item_delete[0])
                    cursor.execute(f'DELETE FROM {table_name} WHERE id = {date_rec[0]}')
                    db.commit()
                    if full_path_picture != '':
                        os.remove(full_path_picture)
                        full_path_picture = ''
                    list_records.delete(ANCHOR)
                    count_record -= 1

                    journal_lab = Label(diary, text=f"Журнал записей (всего: {count_record})", bg='white',
                                        font=('Comic Sans MS', 11))
                    journal_lab.place(x=26, y=125)

                    if count_record == 0:
                        delete_button['state'] = 'disabled'
                        open_button['state'] = 'disabled'
            except IndexError:
                pass
            except FileNotFoundError:
                name.insert(0, 'Изображение не найдено! ')

        def open_record():
            nonlocal full_path_picture, picture_symbol, picture_n
            try:
                delete_trash()
                date_rec = list_records.get(ANCHOR)
                for item4 in cursor.execute(f'SELECT name FROM {table_name} WHERE id = {date_rec[0]}'):
                    name.delete(0, END)
                    n = str(item4)
                    name.insert(0, n[2:-3])
                    name['foreground'] = 'grey20'

                for item5 in cursor.execute(f'SELECT text FROM {table_name} WHERE id = {date_rec[0]}'):
                    text.delete(1.0, END)
                    string = str(item5)
                    num = string.count(r'\n') + 1
                    text.insert(1.0, item5)
                    text.delete(1.0)
                    text.delete(f'{num}.0')

                for item_name_picture in cursor.execute(
                        f'SELECT picture_name FROM {table_name} WHERE id = {date_rec[0]}'):
                    if item_name_picture == ('',):
                        name_label_picture_main['text'] = ''
                    else:
                        picture_n = item_name_picture
                        picture_n = str(picture_n[0])
                        if len(picture_n) <= 24:
                            name_label_picture_main['text'] = picture_n
                        else:
                            name_label_picture_main['text'] = picture_n[:24] + '...'

                for item_path_picture in cursor.execute(
                        f'SELECT picture_path FROM {table_name} WHERE id = {date_rec[0]}'):
                    string_path_picture = str(item_path_picture[0])
                    if item_path_picture != ('',):
                        attached_picture_but['state'] = 'active'
                        full_path_picture = string_path_picture
                        picture_symbol = '✅'
                    else:
                        attached_picture_but['state'] = 'disabled'
                        full_path_picture = ''

            except IndexError:
                pass
            except FileNotFoundError:
                name.insert(0, 'Изображение не найдено! ')

        def new_record():
            try:
                nonlocal full_path_picture, picture_n, picture_symbol
                delete_trash()
                name.delete(0, END)
                name.insert(0, 'no_name')
                text.delete(1.0, END)
                full_path_picture = ''
                picture_n = ''
                picture_symbol = ''
                name_label_picture_main['text'] = picture_n
                attached_picture_but['state'] = 'disabled'
            except FileNotFoundError:
                pass

        def insert_text():
            file = filedialog.askopenfilename(filetypes=[("All files", "*"), ("JPEG", "*.jpg"), ("DOCX", ".docx"),
                                                         ("PNG", "*.png"), ("TXT", "*.txt")])
            try:
                with open(file, 'r', encoding='utf-8') as f:
                    s = f.read()
                    name.delete(0, END)
                    srez = file.rfind('/') + 1
                    name.insert(0, file[srez:-4])
                    text.delete(1.0, END)
                    text.insert(1.0, s)
            except FileNotFoundError:
                pass
            except UnicodeDecodeError:
                if file.find('.docx') != -1:
                    doc = Document(file)
                    text2 = []
                    for paragraph in doc.paragraphs:
                        text2.append(paragraph.text)
                    name.delete(0, END)
                    srez = file.rfind('/') + 1
                    name.insert(0, file[srez:-5])
                    text.delete(1.0, END)
                    text.insert(1.0, '\n'.join(text2))
                elif file.find('.jpeg') != -1 or file.find('.jpg') != -1 or file.find('.png') != -1:
                    photo_level = Toplevel(diary)
                    photo_level.title("Photo")
                    img_open = Image.open(file)
                    w, b = img_open.size
                    w1 = w / 850
                    b1 = b / 520
                    if w1 <= 1 and b1 <= 1:
                        img_open = img_open.resize((w, b), Image.BILINEAR)
                    elif w1 > b1:
                        img_open = img_open.resize((int(w / w1), int(b / w1)), Image.BILINEAR)
                    else:
                        img_open = img_open.resize((int(w / b1), int(b / b1)), Image.BILINEAR)
                    img_tk = ImageTk.PhotoImage(img_open)
                    label_img = Label(photo_level, image=img_tk)
                    label_img.image = img_tk
                    label_img.pack()
                    x_coordinate = (photo_level.winfo_screenwidth() - 850) / 2
                    y_coordinate = (photo_level.winfo_screenheight() - 520) / 2
                    photo_level.wm_geometry("+%d+%d" % (x_coordinate, y_coordinate))
                    photo_level.resizable(False, False)
                    photo_level.mainloop()

        def add_file():
            nonlocal full_path_picture, picture_n, picture_symbol

            def add_picture():
                nonlocal full_path_picture, picture_n, picture_symbol
                try:
                    picture = filedialog.askopenfilename(
                        filetypes=[("All files", "*"), ("JPEG", "*.jpeg"), ("PNG", "*.png")])
                    if picture.find('.jpeg') != -1 or picture.find('.jpg') != -1 or picture.find(
                            '.png') != -1 or picture.find('.ico') != -1 or picture.find('.JPEG') != -1 or \
                            picture.find('.JPG') != -1 or picture.find('.PNG') != -1 or picture.find('.ICO') != -1:
                        picture_n = os.path.basename(picture)
                        path_db = os.path.abspath('db_diary.db')
                        path_db_without_file = path_db[:path_db.rfind('\\') + 1]
                        if not os.path.isdir('.\\pictures'):
                            os.mkdir('pictures')
                        if not os.path.isdir('.\\logs_ps'):
                            os.mkdir('logs_ps')
                        full_path_picture = path_db_without_file + 'logs_ps\\' + picture_n
                        shutil.copyfile(picture, full_path_picture)

                        ok_label_picture['fg'] = 'green2'
                        if len(picture_n) <= 20:
                            name_label_picture['text'] = picture_n
                        else:
                            name_label_picture['text'] = picture_n[:20] + '...'

                        if len(picture_n) <= 24:
                            name_label_picture_main['text'] = picture_n
                        else:
                            name_label_picture_main['text'] = picture_n[:24] + '...'

                        attached_picture_but['state'] = 'active'
                        picture_symbol = '✅'
                        attach_level.focus_force()
                except UnicodeDecodeError:
                    pass
                except FileNotFoundError:
                    pass

            def add_sound():
                pass

            def destroy_attach_window():
                attach_level.destroy()

            attach_level = Toplevel(diary)
            attach_level.title('Attach file')
            x_coordinate = (attach_level.winfo_screenwidth() - 400) / 2
            y_coordinate = (attach_level.winfo_screenheight() - 200) / 2
            attach_level.wm_geometry("+%d+%d" % (x_coordinate, y_coordinate))
            attach_level.geometry('400x200')
            attach_level.resizable(False, False)
            attach_level['bg'] = 'gray90'

            style_add_picture = ttk.Style()
            style_add_picture.configure('add_picture.TButton', font=('Comic Sans MS', 9))
            add_picture = ttk.Button(attach_level, text="Прикрепить картинку", command=add_picture,
                                     style='add_picture.TButton')
            add_picture.place(x=30, y=30)

            style_add_sound = ttk.Style()
            style_add_sound.configure('add_sound.TButton', font=('Comic Sans MS', 9))
            add_sound = ttk.Button(attach_level, text="Записать голосовое сообщение", command=add_sound,
                                   style='add_sound.TButton')
            add_sound.place(x=30, y=70)

            ok_label_picture = Label(attach_level, text='✔', font=('Segoe Script', 12), bg='gray90', fg='gray60')
            ok_label_picture.place(x=174, y=30)

            name_label_picture = Label(attach_level, text='', font=('Comic Sans MS', 10), bg='gray90')
            name_label_picture.place(x=200, y=33)

            ok_label_sound = Label(attach_level, text='✔', font=('Segoe Script', 12), bg='gray90', fg='gray60')
            ok_label_sound.place(x=232, y=70)

            style_ok = ttk.Style()
            style_ok.configure('ok.TButton', font=('Comic Sans MS', 10))
            ok_but = ttk.Button(attach_level, text="OK", command=destroy_attach_window, style='ok_but.TButton')
            ok_but.place(x=302, y=157)

            attach_level.mainloop()

        def send_file():
            pass

        def attached_picture():
            nonlocal full_path_picture, picture_n
            try:
                if full_path_picture.find('.jpg') != -1 or full_path_picture.find('.png') != -1 \
                        or full_path_picture.find('.jpeg') != -1 or full_path_picture.find('.ico') != -1 or \
                        full_path_picture.find('.JPG') != -1 or full_path_picture.find('.PNG') != -1 \
                        or full_path_picture.find('.JPEG') != -1 or full_path_picture.find('.ICO') != -1:
                    photo_level = Toplevel(diary)
                    photo_level.title(picture_n)
                    img_open = Image.open(full_path_picture)
                    w, b = img_open.size
                    w1 = w / 850
                    b1 = b / 520
                    if w1 <= 1 and b1 <= 1:
                        img_open = img_open.resize((w, b), Image.BILINEAR)
                    elif w1 > b1:
                        img_open = img_open.resize((int(w / w1), int(b / w1)), Image.BILINEAR)
                    else:
                        img_open = img_open.resize((int(w / b1), int(b / b1)), Image.BILINEAR)
                    img_tk = ImageTk.PhotoImage(img_open)
                    label_img = Label(photo_level, image=img_tk)
                    label_img.image = img_tk
                    label_img.pack()
                    x_coordinate = (photo_level.winfo_screenwidth() - 850) / 2
                    y_coordinate = (photo_level.winfo_screenheight() - 520) / 2
                    photo_level.wm_geometry("+%d+%d" % (x_coordinate, y_coordinate))
                    photo_level.resizable(False, False)
                    photo_level.mainloop()
            except FileNotFoundError:
                name.insert(0, "Изображение не найдено! ")

        def attached_sound():
            pass

        def search():
            request = google_search.get()
            if request != 'Google':
                if request.find('http') != -1 or request.find('www') != -1:
                    webbrowser.open_new(f"{request}")
                else:
                    webbrowser.open_new(rf"https://www.google.com/search?q={request}")

        def search_enter(_):
            request = google_search.get()
            if request.find('http') != -1 or request.find('www') != -1:
                webbrowser.open_new(f"{request}")
            else:
                webbrowser.open_new(rf"https://www.google.com/search?q={request}")

        def on_entry_click(_):
            if google_search.get() == 'Google':
                google_search.delete(0, "end")  # delete all the text in the entry
                google_search.insert(0, '')  # Insert blank for user input
                google_search.config(foreground='black')

        def on_focusout(_):
            if google_search.get() == '':
                google_search.insert(0, 'Google')
                google_search.config(foreground='grey47')

        def on_entry_click_name(_):
            if name.get() == 'no_name':
                name.delete(0, "end")  # delete all the text in the entry
                name.insert(0, '')  # Insert blank for user input
                name.config(foreground='grey20')

        def on_focusout_name(_):
            if name.get() == '':
                name.insert(0, 'no_name')
                name.config(foreground='grey60')

        user_login = login.get()
        user_password = password.get()
        root.destroy()
        diary = Tk()
        diary.title('My diary')
        x_coord = (diary.winfo_screenwidth() - 920) / 2
        y_coord = (diary.winfo_screenheight() - 560) / 2
        diary.wm_geometry("+%d+%d" % (x_coord, y_coord))
        diary.geometry('920x560')
        diary.resizable(False, False)
        diary['bg'] = 'gray50'
        diary.attributes("-alpha", 0.95)

        img = Image.open('main.jpg')
        imgtk = ImageTk.PhotoImage(img)
        labelq = Label(diary, image=imgtk)
        labelq.image = imgtk
        labelq.pack()

        full_path_picture = ''
        picture_n = ''
        picture_symbol = ''

        full_path_sound = ''
        sound_name = ''
        sound_symbol = ''

        main_label = Label(diary, text='Хлопотное дельце', font=('Segoe Script', 38), bg='gray97')
        main_label.place(relx=0.38, rely=0.06)

        table_name = str(user_login) + str(user_password)
        sql_request = f"""CREATE TABLE IF NOT EXISTS {table_name} (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                login TEXT NOT NULL DEFAULT 1111,
                password TEXT NOT NULL DEFAULT 1111,
                date TEXT,
                name TEXT DEFAULT no_name,
                text TEXT,
                picture_name TEXT,
                picture_path TEXT,
                picture_symbol TEXT,
                sound_name TEXT,
                sound_path TEXT,
                sound_symbol TEXT)"""
        cursor.execute(sql_request)
        db.commit()

        frame_for_records = Frame(diary)
        frame_for_records.place(x=20, y=166)

        list_records = Listbox(frame_for_records, width=36, height=18, bg="gray88", fg='gray10',
                               font=('Comic Sans MS', 9))
        list_records.pack(side="left", fill="y")

        count_record = 0
        for item in cursor.execute(f'SELECT id, date, picture_symbol, name FROM {table_name}'):
            if item[2] == '':
                item1_list = list(item)
                item1_list.remove('')
                item = item1_list
                list_records.insert(0, item)
                count_record += 1
            else:
                list_records.insert(0, item)
                count_record += 1

        scrollbar_right = Scrollbar(frame_for_records, orient=VERTICAL)
        scrollbar_right.pack(side=RIGHT, fill=Y)
        scrollbar_right.config(command=list_records.yview)
        list_records.config(yscrollcommand=scrollbar_right.set)

        scrollbar_down = Scrollbar(frame_for_records, command=list_records.xview, orient=HORIZONTAL)
        scrollbar_down.place(x=0, rely=0.984, relwidth=0.93)
        list_records.config(xscrollcommand=scrollbar_down.set)

        journal_label = Label(diary, text=f"Журнал записей (всего: {count_record})", bg='white',
                              font=('Comic Sans MS', 11))
        journal_label.place(x=26, y=125)

        list_label = Label(diary, text="№                   Дата                  Название", bg='white',
                           font=('Comic Sans MS', 7))
        list_label.place(x=22, y=147)
        '''
        delete_picture_label = Label(diary, text="x", bg='white', font=('Comic Sans MS', 11))
        delete_picture_label.place(x=630, y=185)
        def del_picture():
            delete_picture_label['text'] = ''
        delete_picture_label.bind('Button-1', del_picture)
        '''
        google_search = ttk.Entry(diary, width=35, font=('Comic Sans MS', 8), foreground='gray47')
        google_search.place(x=618, y=121)
        google_search.insert(0, 'Google')
        google_search.bind('<FocusIn>', on_entry_click)
        google_search.bind('<FocusOut>', on_focusout)
        google_search.bind('<Return>', search_enter)

        name_label = Label(diary, text="Название:", bg='white', font=('Comic Sans MS', 11))
        name_label.place(x=315, y=125)

        name = ttk.Entry(diary, width=58, font=('Comic Sans MS', 12), foreground='gray60')
        name.place(x=310, y=150)
        name.insert(0, 'no_name')
        name.bind('<FocusIn>', on_entry_click_name)
        name.bind('<FocusOut>', on_focusout_name)

        text_label = Label(diary, text="Текст:", bg='white', font=('Comic Sans MS', 11))
        text_label.place(x=315, y=185)

        last_size = '12'
        last_font = 'Comic Sans MS'
        text_frame = Frame(diary, width=585, height=284)
        text_frame.place(x=310, y=210)
        text = Text(text_frame, width=58, height=12, bg="gray95", fg='gray15', bd=2,
                    font=(last_font, int(last_size)), wrap=WORD)
        text.focus()
        text.place(x=0, y=0)

        diary.bind_all("<Key>", _on_key_release, "+")

        style_search = ttk.Style()
        style_search.configure('search.TButton', font=('Comic Sans MS', 7), width=3)
        search_but = ttk.Button(diary, text="▶", command=search, style='search.TButton')
        search_but.place(x=868, y=120)

        style_attached_sound = ttk.Style()
        style_attached_sound.configure('attached_sound.TButton', font=('Comic Sans MS', 9), width=3)
        attached_sound_but = ttk.Button(diary, text="+♬", command=attached_sound, style='attached_sound.TButton',
                                        state='disabled')
        attached_sound_but.place(x=405, y=181)

        style_attached_picture = ttk.Style()
        style_attached_picture.configure('attached_picture.TButton', font=('Comic Sans MS', 9), width=3)
        attached_picture_but = ttk.Button(diary, text="+✅", command=attached_picture, style='attached_picture.TButton',
                                          state='disabled')
        attached_picture_but.place(x=370, y=181)

        name_label_picture_main = Label(diary, text=picture_n, font=('Comic Sans MS', 10), bg='white')
        name_label_picture_main.place(x=448, y=186)

        style_send_file = ttk.Style()
        style_send_file.configure('send_file.TButton', font=('Comic Sans MS', 9))
        send_file_but = ttk.Button(diary, text="Отправить ✉", command=send_file, style='send_file.TButton')
        send_file_but.place(x=335, y=506)

        style_addfile = ttk.Style()
        style_addfile.configure('addfile.TButton', font=('Comic Sans MS', 9))
        add_file_but = ttk.Button(diary, text="Прикрепить ➕", command=add_file, style='addfile.TButton')
        add_file_but.place(x=652, y=506)

        style_openfile = ttk.Style()
        style_openfile.configure('openfile.TButton', font=('Comic Sans MS', 9))
        open_file = ttk.Button(diary, text="Открыть файл ☝", command=insert_text, style='openfile.TButton')
        open_file.place(x=531, y=506)

        style_save = ttk.Style()
        style_save.configure('save.TButton', font=('Comic Sans MS', 9), foreground='deepskyblue3')
        save_button = ttk.Button(diary, text='Сохранить ⯆', command=save, cursor="hand2", style='save.TButton')
        save_button.place(x=785, y=506)

        style_setting = ttk.Style()
        style_setting.configure('setting.TButton', font=('Comic Sans MS', 9))
        setting_button = ttk.Button(diary, text='Настройки ⚒', command=setting, style='setting.TButton',
                                    cursor="hand2")
        setting_button.place(x=433, y=506)

        style_save = ttk.Style()
        style_save.configure('delete.TButton', font=('Comic Sans MS', 9), foreground='red')
        delete_button = ttk.Button(diary, text='Удалить ❌', command=delete, cursor="hand2", style='delete.TButton')
        delete_button.place(x=24, y=506)

        style_open = ttk.Style()
        style_open.configure('open.TButton', font=('Comic Sans MS', 9), foreground='green3')
        open_button = ttk.Button(diary, text='Открыть ⏏', command=open_record, cursor="hand2", style='open.TButton')
        open_button.place(x=114, y=506)

        style_new = ttk.Style()
        style_new.configure('new.TButton', font=('Comic Sans MS', 9), foreground='blue')
        new_button = ttk.Button(diary, text='Новый ❒', command=new_record, cursor="hand2", style='new.TButton')
        new_button.place(x=204, y=506)

        if count_record == 0:
            delete_button['state'] = 'disabled'
            open_button['state'] = 'disabled'

        sql_last_user = """CREATE TABLE IF NOT EXISTS last_user (
                    user_id INTEGER PRIMARY KEY AUTOINCREMENT,
                    login TEXT NOT NULL DEFAULT 1111,
                    password TEXT NOT NULL DEFAULT 1111,
                    char TEXT NOT NULL DEFAULT q)"""
        cursor.execute(sql_last_user)
        db.commit()
        install_char = 'q'
        try:
            cursor.execute('SELECT char FROM last_user ORDER BY user_id DESC')
            for item in cursor.fetchone():
                install_char = item
        except TypeError:
            pass
        cursor.execute("INSERT INTO last_user (login, password, char) VALUES (?, ?, ?)",
                       (user_login, user_password, install_char))
        db.commit()

        diary.mainloop()

    # function to registration in diary
    def entrance():
        user_login = login.get()
        user_password = password.get()
        if len(user_login) < 4 or len(user_password) < 4:
            messagebox.showinfo('Внимание', 'Длина логина или пароля меньше 4 символов!')
        else:
            cursor.execute(
                f"SELECT login, password FROM users WHERE login = '{user_login}' AND password = '{user_password}'")
            if cursor.fetchone() is None:
                cursor.execute("INSERT INTO users (login, password) VALUES (?, ?)", (user_login, user_password))
                db.commit()
            main_window()

    def ent(_):
        cursor.execute('SELECT login FROM last_user ORDER BY user_id DESC')
        for item in cursor.fetchone():
            login.insert(1, item)
        cursor.execute('SELECT password FROM last_user ORDER BY user_id DESC')
        for item in cursor.fetchone():
            password.insert(1, item)
        main_window()

    root = Tk()
    root.title('My diary')
    x = (root.winfo_screenwidth() - 920) / 2
    y = (root.winfo_screenheight() - 560) / 2
    root.wm_geometry("+%d+%d" % (x, y))
    root.geometry('920x560')
    root.resizable(False, False)
    root['bg'] = 'gray70'
    root.attributes("-alpha", 0.95)
    enter = 'q'

    with sqlite3.connect('db_diary.db') as db:
        cursor = db.cursor()
        sql = """CREATE TABLE IF NOT EXISTS users (
            user_id INTEGER PRIMARY KEY AUTOINCREMENT,
            login TEXT NOT NULL DEFAULT 1111,
            password TEXT NOT NULL DEFAULT 1111)"""
        cursor.execute(sql)
        db.commit()

    login_frame = LabelFrame(root, padx=10, pady=10, text='Вход', bg='gray70', font=('Comic Sans MS', 11))
    login_frame.pack(padx=0, pady=208)
    Label(login_frame, text="Логин", background='gray70', font=('Comic Sans MS', 9)).grid(row=0)
    Label(login_frame, text="Пароль", bg='gray70', font=('Comic Sans MS', 9), pady=7).grid(row=1)
    login = ttk.Entry(login_frame, font=('Comic Sans MS', 11))
    login.grid(row=0, column=1, columnspan=2, sticky=W)
    password = ttk.Entry(login_frame, font=('Comic Sans MS', 11), show='*')
    password.grid(row=1, column=1, columnspan=2, sticky=W)
    style = ttk.Style()
    style.configure('enter.TButton', foreground='green3', font=('Comic Sans MS', 9))
    ttk.Button(login_frame, text='Войти', command=entrance, style='enter.TButton',
               width=8, cursor="hand2").grid(row=2, column=2, sticky=W)

    img_phon = Image.open('12.jpg')
    imgtk_phon = ImageTk.PhotoImage(img_phon)
    label = Label(root, image=imgtk_phon)
    label.image = imgtk_phon
    label.pack()

    try:
        cursor.execute('SELECT char FROM last_user ORDER BY user_id DESC')
        for i in cursor.fetchone():
            enter = i
        root.bind(f'<Control-{enter}>', ent)
    except sqlite3.OperationalError:
        pass

    root.mainloop()


if __name__ == '__main__':
    main()
